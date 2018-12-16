# -*- coding=utf-8 -*-
#创建并写入word文档
import os
import docx
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.shared import Inches

templetpath = os.path.dirname(__file__)+'/word.docx'
doc = None
numR = 0
numC = 0
a=0
def CreatTable(mRows,mCols,strname,ftsize=7.5):   #创建表格
    global  doc
    global a
    #doc.add_heading(strname, level=0)
    #doc.add_paragraph("strname")
    if not strname:
        pass
    else:
        p = None
        if a == 0:
            p = doc.paragraphs[0]
        else:
            p = doc.add_paragraph()
        if p:
            run = p.add_run(strname)
            #run.font.size = Pt(15)
            run.font.size = Pt(ftsize)
            run.bold = True
    #global numC, numR    #数据当前进入表格的行列数
    #numC=numR=0
    a=a+1
    table=doc.add_table(rows=mRows,cols=mCols)
    #SetTableStyle(table)
    #str = "------------------------------------"
    #doc.add_paragraph(str * 3)

    return a-1

def Blank(num):
    global doc
    for i in range(0,num):
        p = doc.add_paragraph()


def PullValue(x,y,str,id=-1):   #存数据进入表格
    #global  numC, numR
    global  a
    if a==0:        #判断是否没有生成表
        return -1
    if id==-1 or id>a:
        id=a-1
    #获取表格的行列数
    rsum = len(doc.tables[id].rows)    #索引表格的总行数
    csum = len(doc.tables[id].columns)      #索引表格的总列数
    if x>=rsum or y>=csum or x<0 or y<0:
        return -1
    doc.tables[id].rows[x].cells[y].text = str
    doc.tables[id].rows[x].cells[y].vertical_alignment = 1
    return 0

def saveword(path):
    #doc.save( os.path.dirname(__file__)+"/Result.docx")
	doc.save(path)

def gettable(id):
    ncount = len(doc.tables)
    if ncount == 0:
        return None
    if id == -1 or id > ncount:
        id = ncount - 1
    return doc.tables[id]

def SetColWidth(icol,width,id=-1):
    curtbl = gettable(id)
    if curtbl == None:
        return -1

    rsum = len(curtbl.rows)    #索引表格的总行数
    csum = len(curtbl.columns)      #索引表格的总列数
    if icol < 0 or icol >= csum:
        return -1

    #curtbl.autofit = False
    curtbl.columns[icol].width = Inches(width)
    
    curtbl.autofit = False
    tblGrid = curtbl._tbl.tblGrid
    gridCol = tblGrid[icol]
    gridCol.w = Inches(width)
    for tr in curtbl._tbl.tr_lst:
        tc = tr.tc_at_grid_col(icol)
        tc.width = Inches(width)
    return curtbl.columns[icol].width

def SetTableStyle(curtbl):
    if not curtbl:
        return

    #curtbl.style = 'Table Grid'
    tblPr = curtbl._tblPr
    #tblBorders = tblPr.get_or_add_tblBorders()
    tblBorders = OxmlElement('w:tblBorders')
    lside = ['w:top','w:left','w:bottom','w:right','w:insideH','w:insideV']
    lcontent = [('w:color','auto'),('w:space','0'),('w:sz','4'),('w:val','dotted')]
    #lcontent = lcontent[::-1]
    print(lcontent)
    for side in lside:
        tmp = OxmlElement(side)
        for content in lcontent:
            tmp.set(qn(content[0]),content[1])
        tblBorders.append(tmp)
    tblPr.append(tblBorders)
    #print(tblPr)
    #print(tblBorders)
    #top = OxmlElement('w:top')
    #top.set(qn('w:color'),'auto')


def SetCellTextStyle(irow,icol,underline=-1,alignment=-1,id=-1):
    curtbl = gettable(id)
    if curtbl == None:
        return
    cell = curtbl.cell(irow,icol)
    p = cell.paragraphs[0]
    p.runs[0].underline = None if underline == -1 else underline
    p.alignment = None if alignment == -1 else alignment
    #p.paragraph_format.left_indent = Inches(0.5)

def MergeCell(irow1,icol1,irow2,icol2,id=-1):
    curtbl = gettable(id)
    if curtbl == None:
        return -1
    cell1 = curtbl.cell(irow1,icol1)
    cell2 = curtbl.cell(irow2,icol2)
    cell1.merge(cell2)
    ret = abs(icol2-icol1) if irow1 == irow2 else abs(irow2 - irow1)
    return ret

def GetRowWidth(id=-1):
    curtbl = gettable(id)
    if curtbl == None:
        return -1
    width = 0
    for col in curtbl.columns:
        width += col.width
    return width

#创建内存中的word文档对象
def Newdocx():
    global  doc
    doc = docx.Document(templetpath)
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size=Pt(7.5)
    global a
    a = 0
    
def text():
    Newdocx()
    CreatTable(4,3,"功能段")
    PullValue(1,2,"材料")
    Blank(6)
    saveword("C:\\Users\\SOFT\\Desktop\\text.docx")

#text()
